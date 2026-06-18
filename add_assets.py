from docx import Document
import glob,os,re

desktop=os.path.expanduser('~/Desktop')
files=[f for f in glob.glob(desktop+'/*V6*') if not os.path.basename(f).startswith('~$')]
d=Document(files[0])
body=d.element.body
ns='http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# Asset definitions per segment
assets={
'1A': [
    '【本段素材】',
    '表情:F1(专注) | 场景:男生房间 | 道具:RGB键盘·显示器(暗黑UI)·能量饮料罐',
    '光源:RGB蓝紫背光+屏幕白光 | 首帧图:镜A1·手与键盘近景',
],
'1B': [
    '【本段素材】',
    '表情:F1(平静) | 场景:同上 | 道具:电竞椅',
    '光源:屏幕金光+投影微光 | 首帧图:镜B4·伸懒腰中景',
],
'1C': [
    '【本段素材】',
    '表情:F1(疲惫) | 场景:同上 | 道具:投影时钟·空饮料罐 | Q版:Q1打呵欠(静态图)',
    '光源:投影白光(房间暗) | 首帧图:镜C7·投影时钟大特写',
],
'2A': [
    '【本段素材】',
    '表情:F2(不安) | 场景:暗紫色三维隧道·发光几何碎片(三角/六边)·白色亮光',
    '光源:隧道暗紫+白色亮光 | 首帧图:镜A1·纯黑浮现侧脸',
],
'2B': [
    '【本段素材】',
    '表情:F2(恐惧→接受) | 场景:同上隧道 | 特效:暖金色光雾',
    '光源:隧道暗紫+暖金光 | 首帧图:镜B4·面部紧绷近景',
],
'2C': [
    '【本段素材】',
    '表情:F2(警觉) | 场景:灰色碎石地面·浓雾·石路轮廓 | Q版:Q2拍脸震惊(静态图)',
    '光源:冷灰散射光 | 首帧图:镜C7·眼皮睁开大特写',
],
'3A': [
    '【本段素材】',
    '表情:F3(谨慎) | 场景:同上·浓雾碎石地面',
    '光源:灰白散射 | 首帧图:镜A1·戒备松拳近景',
],
'3B': [
    '【本段素材】',
    '表情:F3(犹豫→决定) | 场景:同上·小径轮廓·石门远景',
    '光源:同上 | 首帧图:镜B4·发现小径近景',
],
'3C': [
    '【本段素材】',
    '表情:F3(惊恐) | 场景:同上·E1石门+紫色光晕 | Q版:Q3好奇震惊(静态PNG)',
    '光源:紫色光晕映照 | 首帧图:镜C7·快步接近石门中景',
],
'4A': [
    '【本段素材】',
    '表情:F4(眩晕→震撼) | 场景:E2深坑全景(已有场景可复用)·巨剑(右后方斜插)',
    '光源:琥珀色天空(顶光)+坑底蓝光(底光) | 首帧图:镜A1·紫色碎裂跌落',
],
'4B': [
    '【本段素材】',
    '表情:F4(震撼→凝视) | 场景:同上·巨剑前景 | 光照:同上',
    '首帧图:镜B4·深坑全景(男生POV)',
],
'4C': [
    '【本段素材】',
    '表情:F4(凝视→确认) | 场景:E5石碑+异世界文字·巨剑(右后方) | Q版:Q4确认震惊(静态PNG)·Q6「不是梦」(静态图+配音)',
    '光照:同上 | 首帧图:镜C7·走下斜坡中景',
],
'5A': [
    '【本段素材】',
    '表情:F5(吸引) | 场景:E3石台+刀(插台中)·巨剑(远处右后方) | 道具:刀三视图',
    '特效:光羽V1(刀身浮现) | 光照:琥珀天+蓝底光 | 首帧图:镜A1·接近石台近景',
],
'5B': [
    '【本段素材】',
    '表情:F5(警觉→恐惧) | 场景:同上·石台碎裂废墟 | 怪物:一阶段(已有) | 特效:灵魂之火(已有)',
    '光照:同上+灵魂之火蓝白 | 首帧图:镜B4·涌烟中景',
],
'5C': [
    '【本段素材】',
    '表情:F5(恐惧→拼命→无奈) | 场景:同上·E4雾墙 | 怪物:一阶段(同上) | 特效:白光光环V2',
    'Q版:Q5抓狂(已有·静态图) | 光照:同上 | 首帧图:镜C7·转身逃跑中景',
],
'5D': [
    '【本段素材】',
    '表情:F5(觉悟→决意) | 场景:同上·雾墙(身后)·巨剑(右后方) | 怪物:一阶段(同上)',
    '特效:光羽V1(重新浮现) | 光照:怪物白光(70%)+蓝底光(30%) | 首帧图:镜D10·低头看刀近景',
],
'6A': [
    '【本段素材】',
    '表情:F6(紧张) | 场景:深坑·雾墙(身后)·巨剑(右后方)·地面碎裂 | 怪物:一阶段(同上)',
    '特效:光羽V1·白光光环V2 | 光照:琥珀天+蓝底光+怪物白光 | 首帧图:镜F1·触手砸地',
],
'6B': [
    '【本段素材】',
    '表情:F6(燃) | 场景:同上·飞行巨石 | 怪物:一阶段(同上) | 特效:光羽V1·白光光环V2',
    '光照:同上 | 首帧图:镜F5·多触手砸地广角',
],
'6C': [
    '【本段素材】',
    '表情:F6(喘息·脸伤持续) | 场景:同上·石碑(E5)·雾墙·巨剑(右后方) | 怪物:二阶段(已有)',
    '特效:光羽V1·白光光环V2·V3结尾卡 | 光照:琥珀天灰化+蓝底光20% | 首帧图:镜F9·白光渐散落地',
],
}

# Find each segment's end (─── before next segment) and insert assets
seg_order=['1A','1B','1C','2A','2B','2C','3A','3B','3C','4A','4B','4C','5A','5B','5C','5D','6A','6B','6C']
seg_names={'1A':'第一段 A','1B':'第一段 B','1C':'第一段 C',
           '2A':'第二段 A','2B':'第二段 B','2C':'第二段 C',
           '3A':'第三段 A','3B':'第三段 B','3C':'第三段 C',
           '4A':'第四段 A','4B':'第四段 B','4C':'第四段 C',
           '5A':'第五段 A','5B':'第五段 B','5C':'第五段 C','5D':'第五段 D',
           '6A':'第六段 A','6B':'第六段 B','6C':'第六段 C'}

import lxml.etree as ET

inserted=0
for seg in seg_order:
    seg_title=seg_names[seg]
    my_assets=assets[seg]
    
    # Find this segment's heading paragraph index
    seg_start=None
    for i,p in enumerate(d.paragraphs):
        if seg_title in p.text and '|' in p.text and '镜' in p.text:
            seg_start=i; break
    
    if seg_start is None: continue
    
    # Find the next segment's heading (to know where this segment ends)
    next_start=None
    for j in range(seg_start+1, len(d.paragraphs)):
        for s2 in seg_order:
            if seg_names[s2] in d.paragraphs[j].text and '|' in d.paragraphs[j].text and '镜' in d.paragraphs[j].text:
                next_start=j; break
        if next_start: break
    
    if next_start is None: next_start=len(d.paragraphs)
    
    # Find the last ─── separator before next_start (that's the end of this segment)
    insert_pos=None
    for j in range(next_start-1, seg_start, -1):
        if '───' in d.paragraphs[j].text.strip():
            insert_pos=j; break
    
    if insert_pos is None: continue
    
    # Insert asset lines right before this ───
    pe=d.paragraphs[insert_pos]._element
    parent=pe.getparent()
    pi=list(parent).index(pe)
    
    for line in reversed(my_assets):
        np=ET.SubElement(parent,'{'+ns+'}p')
        nr=ET.SubElement(np,'{'+ns+'}r')
        nt=ET.SubElement(nr,'{'+ns+'}t')
        nt.text=line
        nt.set('{http://www.w3.org/XML/1998/namespace}space','preserve')
        parent.insert(pi,np)
    inserted+=1

d.save(files[0])
print(f'Inserted assets for {inserted}/{len(seg_order)} segments')
