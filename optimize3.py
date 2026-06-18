from docx import Document

doc = Document(r'C:\Users\lintianhao\Desktop\分镜稿.docx')
c = 0

for p in doc.paragraphs:
    t = p.text
    if not t: continue

    # CPY2: 四段镜1 break up triple dashes
    if '先是一瞬间的绝对寂静——只有细碎的风吹过碎石地面的微尘在飘动——然后' in t:
        for r in p.runs:
            if '先是一瞬间的绝对寂静' in r.text:
                r.text = r.text.replace(
                    '先是一瞬间的绝对寂静——只有细碎的风吹过碎石地面的微尘在飘动——然后',
                    '先是一瞬间的绝对寂静，细碎的风吹过碎石地面的微尘在飘动。然后'
                )
                c += 1; break

    # CON1: change "已见底" to "还剩最后一口"
    if '能量饮料罐已见底' in t:
        for r in p.runs:
            if '已见底' in r.text:
                r.text = r.text.replace('已见底', '还剩最后一口')
                c += 1; break

    # CON2: remove "眼角挂一滴泪珠"
    if '张嘴打呵欠 + 眼角挂一滴泪珠' in t:
        for r in p.runs:
            if '眼角挂一滴泪珠' in r.text:
                r.text = r.text.replace('张嘴打呵欠 + 眼角挂一滴泪珠', '张嘴打呵欠')
                c += 1; break

    # LOG1: "accelerates again" → "maintains his speed and pushes even harder"
    if '男生再次加速' in t and '音爆' in t:
        for r in p.runs:
            if '再次加速' in r.text:
                r.text = r.text.replace('再次加速', '保持高速并全力推进')
                c += 1; break

    # NAR1: add F12 after F11
    if '─── 未完待续 ───' in t:
        # Insert F12 before this line
        for r in p.runs:
            if '未完待续' in r.text:
                r.text = r.text.replace(
                    '─── 未完待续 ───',
                    ''
                )
                c += 1; break

# Add F12 and new ending after F11
f12_lines = [
    '',
    '【F12】二阶段第一击',
    'Before the transformation fully completes, the second-phase monster lashes out with a single devastating strike — a massive jagged tendril of black-tinged fog shoots directly toward the camera. The boy barely raises his blade in time to block. The impact sends him skidding backward across the shattered ground, feet digging twin trenches through the stone, sparks erupting from the blade. He holds his ground but the blade is trembling. Close-up on his eyes — still calculating, still alive. The monster rises to its full second-phase height, dwarfing everything. Cut to black.',
    '',
    '─── 未完待续 ───',
]

for line in f12_lines:
    doc.add_paragraph(line)

doc.save(r'C:\Users\lintianhao\Desktop\分镜稿.docx')
print(f'Changes: {c}')
