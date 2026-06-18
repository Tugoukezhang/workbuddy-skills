from docx import Document
doc = Document(r'C:\Users\lintianhao\Desktop\提示词2.docx')
with open(r'C:\Users\lintianhao\WorkBuddy\2026-06-12-10-12-10\hint2.txt', 'w', encoding='utf-8') as f:
    for p in doc.paragraphs:
        f.write(p.text + '\n')
    f.write('=== TABLES ===\n')
    for t in doc.tables:
        for row in t.rows:
            cells = [c.text for c in row.cells]
            f.write(' | '.join(cells) + '\n')
print('Done')
