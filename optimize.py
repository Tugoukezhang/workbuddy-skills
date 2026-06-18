from docx import Document
from docx.shared import Pt

doc = Document(r'C:\Users\lintianhao\Desktop\分镜稿.docx')

changes = 0

for i, p in enumerate(doc.paragraphs):
    txt = p.text

    # SW1: 五段镜17B - add subtle smile at end of confrontation
    if '身后的雾墙纹丝不动——他无路可退' in txt:
        new_txt = txt.replace(
            '身后的雾墙纹丝不动——他无路可退。',
            '身后的雾墙纹丝不动——他无路可退。右嘴角几乎不可察觉地微微上扬——和打BOSS获胜时一模一样的弧度，一闪即逝。'
        )
        if new_txt != txt:
            for run in p.runs:
                if '身后的雾墙纹丝不动' in run.text:
                    run.text = run.text.replace(
                        '身后的雾墙纹丝不动——他无路可退。',
                        '身后的雾墙纹丝不动——他无路可退。右嘴角几乎不可察觉地微微上扬——和打BOSS获胜时一模一样的弧度，一闪即逝。'
                    )
                    changes += 1
                    break
            else:
                if p.runs:
                    p.runs[-1].text = p.runs[-1].text + '右嘴角几乎不可察觉地微微上扬——和打BOSS获胜时一模一样的弧度，一闪即逝。'
                    changes += 1

    # AD2 + SW2: 四段镜1前 - add quiet transition moment
    if '画面从全紫光中碎裂展开' in txt:
        new_txt = txt.replace(
            '画面从全紫光中碎裂展开。',
            '画面从全紫光中碎裂展开。先是一瞬间的绝对寂静——只有细碎的风吹过碎石地面的微尘在飘动——然后，紫色光晕在画面中炸裂成碎片。'
        )
        if new_txt != txt:
            for run in p.runs:
                if '画面从全紫光中碎裂展开' in run.text:
                    run.text = run.text.replace(
                        '画面从全紫光中碎裂展开。',
                        '画面从全紫光中碎裂展开。先是一瞬间的绝对寂静——只有细碎的风吹过碎石地面的微尘在飘动——然后，紫色光晕在画面中炸裂成碎片。'
                    )
                    changes += 1
                    break

    # IP1: 五段镜2 - knife hilt has same script symbols as tablets
    if '石台中央斜插着一把精美的长刀' in txt and '刀身映出坑底蓝光的微光' in txt:
        new_txt = txt.replace(
            '刀身映出坑底蓝光的微光。',
            '刀身映出坑底蓝光的微光。刀柄末端刻有一圈极细微的符号——和石碑上那行文字的笔画系统如出一辙。'
        )
        if new_txt != txt:
            for run in p.runs:
                if '刀身映出坑底蓝光的微光' in run.text:
                    run.text = run.text.replace(
                        '刀身映出坑底蓝光的微光。',
                        '刀身映出坑底蓝光的微光。刀柄末端刻有一圈极细微的符号——和石碑上那行文字的笔画系统如出一辙。'
                    )
                    changes += 1
                    break

    # GAME3: F10 - add hesitation then resolve
    if 'he is still alive, he is still fighting' in txt and 'He stands next to' in txt:
        new_txt = txt.replace(
            'he is still alive, he is still fighting.',
            'he is still alive, he is still fighting. For a split second his grip loosens — a flicker of doubt — then tightens again, harder than before.'
        )
        if new_txt != txt:
            for run in p.runs:
                if 'he is still fighting' in run.text:
                    run.text = run.text.replace(
                        'he is still alive, he is still fighting.',
                        'he is still alive, he is still fighting. For a split second his grip loosens — a flicker of doubt — then tightens again, harder than before.'
                    )
                    changes += 1
                    break

    # POST3: F11 end - add transition interface to segment 7
    if 'Clean bold anime composition. To be continued.' in txt and 'Phase two transformation' in txt:
        new_txt = txt.replace(
            'To be continued.',
            'To be continued. Fade to black. In the darkness, the silhouette of the transformed second-phase monster flickers for a single frame — then black.'
        )
        if new_txt != txt:
            for run in p.runs:
                if 'To be continued' in run.text:
                    run.text = run.text.replace(
                        'To be continued.',
                        'To be continued. Fade to black. In the darkness, the silhouette of the transformed second-phase monster flickers for a single frame — then black.'
                    )
                    changes += 1
                    break

    # AUD3-Q8: 三段镜4 and 镜5 - shorten the hesitation sequence
    # Remove the second foot shuffle in 镜4
    if '碎石在脚底两次碾动' in txt and '大特写犹豫' in doc.paragraphs[i-2].text if i >= 2 else False:
        for run in p.runs:
            if '碎石在脚底两次碾动' in run.text:
                run.text = run.text.replace('碎石在脚底两次碾动', '碎石在脚底碾动')
                changes += 1
                break
        else:
            p.text = p.text.replace('碎石在脚底两次碾动', '碎石在脚底碾动')
            changes += 1

doc.save(r'C:\Users\lintianhao\Desktop\分镜稿.docx')
print(f'Done. {changes} changes applied.')
