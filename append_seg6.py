from docx import Document
from docx.shared import Pt
import copy

doc = Document(r'C:\Users\lintianhao\Desktop\分镜稿.docx')

# Get Normal style for reference
normal_style = doc.styles['Normal']

# ===== SIXTH SEGMENT CONTENT =====
lines = [
    "",
    "╔══════════════════════════════════════════════════════════════╗",
    "║                      第六段 · 初战                              ║",
    "║   11镜 ·  触手砸地 → 闪避 → 反击 → 跳斩 → 暴击 → 对峙 → 二阶段   ║",
    "╚══════════════════════════════════════════════════════════════╝",
    "",
    "前情：五段镜17B，男生双手持刀对峙怪物，雾墙在身后，画面定格。",
    "",
    "环境：深坑遗迹，灰褐琥珀色天空，坑底蓝光脉动，雾墙在身后。",
    "地面统一为\"cracked and shattered battlefield ground\"。",
    "",
    "格式说明：英文 Prompt，框架式写法。每镜一个核心动作，不写连续动作。",
    "镜头和人物同时动。AI 自己分配节奏。后期速度/转场/BGM在剪辑台完成。",
    "战斗不用Q版。",
    "",
    "道具：男生+刀+光羽+手腕白光光环 | 怪物+雾触手(圆柱形+3-4细雾丝)+头部双光核",
    "二阶段怪物(已出图，吸取石碑灵魂后进化) | 斜插地面巨剑(剑刃缺了大口) | 散落石碑(四段同一批)",
    "",
    "【F1】怪物触手砸地",
    "A massive fog tentacle smashes down onto the shattered stone ground from above, cracking the surface and spraying broken rocks and dust outward. A visible shockwave ring radiates across the ground. The young boy in white T-shirt stands at the edge of the impact, gripping his long blade with both hands, his hair blowing back from the shockwave. Low-angle ground-level camera, strong motion blur on debris, deep blue light pulsing from the pit bottom. Cinematic anime action shot.",
    "",
    "【F2】侧身闪避·触手擦镜",
    "Another tentacle sweeps in horizontally from the side, passing extremely close to the camera lens with heavy motion blur, filling the foreground out of focus. Focus pull on the boy in the midground — he leans sideways and slides away from the attack in sharp focus, body low and balanced, blade held behind him with white feather-light particles trailing. The fog wall behind him glows faintly from the reflected light. Fast dynamic tracking shot, strong sense of speed. Clean anime composition.",
    "",
    "【F3】冲刺反击",
    "The boy suddenly dashes forward across the cracked battlefield ground toward the monster. Low-angle camera tracks backward rapidly, his wrist glows with a white ring of light, speed lines streak across the frame, broken rocks kick up behind his footsteps. Blade held low and ready, white energy trailing from the edge. He plants his front foot hard, pivots his hips, and starts a horizontal swing. Strong sense of acceleration, motion blur, dynamic cinematic anime shot.",
    "",
    "【F4】刀光横斩触手",
    "The boy completes the horizontal slash, cutting across a low-hanging fog tendril at chest height. The blade cuts through with a burst of white feather-like light particles exploding outward. The fog tears open along the slash line before slowly closing. Sparks fly where the edge meets the mist. The reflected white light from the slash briefly illuminates the fog wall behind him, revealing a faint, unclear silhouette within the mist for a split second. Low camera angle, strong impact frame. Cinematic anime action.",
    "",
    "【F5】怪物多触手砸地·巨石崩起",
    "The monster roars silently and slams multiple thick fog tentacles down onto the ground simultaneously. The stone surface cracks open violently, huge chunks of rock are launched upward into the air, dust and debris fill the frame. The boy braces himself, shielding his face with one arm, blade ready. A flying rock shard cuts a thin bleeding line across his cheek. Wide-angle shot, massive destruction, blue pit light flickering through the chaos. Intense cinematic atmosphere.",
    "",
    "【F6】跳上飞行巨石",
    "Amidst the chaos, his eyes lock onto a massive stone slab flying upward — the slab soars up to the monster's shoulder height. He crouches, then leaps upward toward it. One foot lands first on the rough broken edge of the flying rock, his body tilted forward, second leg still extended behind, using it as a stepping platform mid-air. His white ring of light flares brightly on his wrist, legs coiled for the next jump. Camera tilts upward following the trajectory, debris floating in slow motion around him. Dramatic anime action composition.",
    "",
    "【F7】空中跳斩·直取光核",
    "The boy launches off the stone slab mid-air, soaring toward the monster's head, both hands gripping the blade overhead, the weapon glowing brilliantly with white feather lights trailing behind. The camera tilts up sharply from below, capturing him against the amber sky, blade poised directly at one of the monster's twin glowing blue-white core eyes. Extreme sense of height and speed, dynamic perspective. Peak cinematic anime shot.",
    "",
    "【F8】光核暴击·蓝白共鸣",
    "The blade strikes one of the monster's core eyes directly. A massive burst of white light explodes from the impact point first, then the blue light at the bottom of the pit surges violently in response, white and blue energy briefly intertwine across the frame. The monster's entire fog body convulses and staggers. The screen goes white briefly during the fall. Low-angle upward shot. Extreme brightness contrast. Climactic anime impact shot.",
    "",
    "【F9】慢镜落地",
    "The boy lands on the cracked and shattered battlefield ground, one knee bent, blade tip touching the ground beside him for balance. A brief pulse of white light from his wrist absorbs the fall impact. A circular shockwave of dust and small rocks radiates outward from his landing point. Slow-motion shot, his hair settles. He glances at the thin bleeding cut on his cheek, then ignores it — his focus already back on the enemy. White feather particles drifting down around him like soft snow. Low-angle camera, calm after the climax. Cinematic anime resolution frame.",
    "",
    "【F10】对峙呼吸·巨剑入画",
    "The boy slowly straightens up, breathing hard but controlled. A flicker of realization crosses his face — he is still alive, he is still fighting. He stands next to the giant broken sword stuck in the ground, its blade chipped and scarred with a large notch. The giant sword and the boy form a diagonal line across the frame. Distant broken stone tablets visible behind the monster. The blade in his hand still glimmers with faint white light. The monster in the distance staggers but does not fall. The blue light from the pit pulses faster than before. The amber sky barely perceptibly beginning to dim. Standoff composition, strong silhouette contrast, tense quiet before the next storm.",
    "",
    "【F11】二阶段钩子",
    "The staggered monster reaches out with its fog tendrils toward the broken stone tablets scattered around the pit. The stone tablets tremble and crack slightly. Ghostly blue-white soul energy flows from the ancient carved tablets into the monster's body. The boy tenses forward instinctively, then halts — the absorption is too fast. Its fog form twists and reshapes, growing larger, more defined, new jagged forms emerging within the mist. The twin core eyes flash erratically, then stabilize into a brighter, more menacing glow. The white feather lights on the boy's blade flicker and lean toward the absorbing monster, as if recognizing the soul energy. The boy stands blade raised, eyes narrowing, calculating, the thin cut on his cheek still bleeding. Darker atmosphere, the amber sky dimming, blue pit light pulsing violently. Wide shot, low angle. Phase two transformation. Clean bold anime composition. To be continued.",
    "",
    "─── 未完待续 ───",
]

for text in lines:
    p = doc.add_paragraph()
    if text == "":
        p.text = ""
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
    else:
        p.text = text

doc.save(r'C:\Users\lintianhao\Desktop\分镜稿.docx')
print(f'Done. Total paragraphs: {len(doc.paragraphs)}')
