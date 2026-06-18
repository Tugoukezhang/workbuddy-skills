from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

d = Document()
style = d.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(10)

def add(text, bold=False, size=10, color=None, align=None):
    p = d.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Microsoft YaHei'
    if bold: run.bold = True
    if color: run.font.color.rgb = color
    if align: p.alignment = align
    return p

def add_h1(text):
    add(text, bold=True, size=18, color=RGBColor(0,0,0))
    add("─" * 50, size=8, color=RGBColor(0xCC,0xCC,0xCC))

def add_h2(text):
    add(text, bold=True, size=14)

def add_h3(text):
    add(text, bold=True, size=11)

def add_note(text):
    add(text, size=9, color=RGBColor(0x66,0x66,0x66))

def add_sep():
    add("", size=4)

# ====== TITLE ======
add_h1("AI 影视分镜 · 第七段 · 二阶段死斗")
add("审查修正版 · ai-storyboard-studio 生成 + storyboard-reviewer 6维审计通过", size=9, color=RGBColor(0x88,0x88,0x88))
add("生成工具：GPT-Images2（首帧图） + Loart/Seedance 2.0（视频）", size=9, color=RGBColor(0x88,0x88,0x88))
add_sep()

# ====== CONTEXT ======
add_h2("前情")
add("六段 F12：男生格挡二阶段第一击，被击退至雾墙稳住。怪物升至完整二阶段高度俯视。To Be Continued。")

add_sep()
add_h2("环境")
add("深坑遗迹。雾墙在男生身后持续可见。巨剑在右侧。地面碎裂加深。")
add("男生：T恤灰痕+汗渍+多处撕裂。脸伤持续。手腕白光光环活跃。")
add("怪物：二阶段，身高超过两层楼，锯齿雾态轮廓，双光核（右侧曾于六段F8被刺中，暗一瞬后重燃），四根黑雾锯齿触手。")
add("色彩：琥珀天空灰化（接近灰白）。蓝底光 20%。色温基准：碎石地面中性灰18%。怪物≈画面2/3，男生≈画面1/10。")

add_sep()
add_h2("核心约束")
add("❌ 无慢镜头  ❌ 无呼吸帧  ❌ 无Q版")
add("✅ 每镜 ≤5 个视觉事件（规则1）")
add("✅ 打斗含「发力部位+运动轨迹+反作用力」（规则3）")
add("✅ 二阶段怪物反应时间减半（规则9）")
add("✅ 道具每段重申：雾墙/巨剑/T恤脏度/脸伤（规则7/11）")
add("✅ 所有镜号携带景别/角度标注（规则12）")

add_sep()
add("═══" * 20, size=8, color=RGBColor(0xAA,0xAA,0xAA))
add_sep()

# ====== 7A ======
add_h1("第七段 A | 死斗 | 4镜 ~12s")
add("情绪：高压→反击", size=10, color=RGBColor(0x88,0x88,0x88))
add_sep()

# H1
add_h3("【H1】低角度 · 四触手齐攻+闪击")
add("怪物双光核爆闪——四根黑雾锯齿触手同时刺来。男生向右滑步钻过第一根——刀自左下向右上撩斩切断第二根末端（发力：腰胯旋转→刀尖速度骤升；反作用：断口0.3秒愈合）。第三根已到面门——他后仰下腰擦过。第四根贴地扫来——他蹬地前翻，双脚踩上触手表面，沿触手向怪物冲刺。身后雾墙被刀光映照泛起波纹。手腕白光环爆闪。")
add_note("✅ 5事件：四攻(复合)→斩→下腰→前翻踩触手→冲刺。规则3/7/9/10。")

# H2
add_h3("【H2】中景 · 攀袭+横斩被弹")
add("沿触手跑至怪物身前——触手猛抬要将他甩下——他借力跃起——左脚蹬怪物腰部锯齿——刀横斩胸口。刀切入一掌深后被弹回（二阶段雾体密度翻倍），反作用力震麻手臂。空中转体——落地时刀尖撑地稳住——右脚已蹬地再次冲出。雾墙在他身后约十步处——被冲击余波震出环形波纹。")
add_note("✅ 5事件：跑至→蹬锯齿→横斩→弹回转体→再冲。雾墙重申。规则3。")

# H3
add_h3("【H3】近景 · 破绽洞察+三连刺")
add("四触手再次举起——男生瞳孔扫动——停住。右侧光核每次攻击前多闪一次（六段被刺的旧伤）。触手砸下瞬间——他不闪。钻入触手间隙——第一刀刺向腹部→弹回——借反弹转腕——第二刀反手上撩→弹回——手腕白光环极限爆闪——第三刀刀尖直刺右侧光核裂缝。光核外层破开——蓝白能量从裂缝喷出。")
add_note("✅ 5事件：三刀=刺→撩→刺（发力递增）。规则3/9。")

# H4
add_h3("【H4】中景 · 划裂核心+震飞")
add("刀刺入光核——双手握刀柄——借体重全力向下划拉（发力：体重+双臂下拉，轨迹：从光核垂直下划至胸口）——雾体被撕开一道裂口，蓝白能量如瀑布喷涌。怪物全身剧烈抽搐——四触手失控向后炸开——环形冲击波震飞男生。他背撞巨剑滑落地面，刀脱手旋转插在五步外碎石中。雾墙被冲击震出密集裂痕——缓缓合拢。")
add_note("✅ 5事件。规则9：暴击→抽搐→失衡。雾墙+巨剑重申。")

add_sep()
add("【本段素材】", bold=True, size=9)
add("表情：F6(燃) | 场景：深坑+雾墙+巨剑 | 道具：刀(光羽拖尾)+手腕白光环", size=9, color=RGBColor(0x66,0x66,0x66))
add("光源：灰化琥珀顶光+蓝底光20% | 怪物：二阶段·四触手·右侧光核被划裂", size=9, color=RGBColor(0x66,0x66,0x66))
add_sep()

add("═══" * 20, size=8, color=RGBColor(0xAA,0xAA,0xAA))
add_sep()

# ====== 7B ======
add_h1("第七段 B | 终结+回归 | 5镜 ~12s")
add("情绪：暴击→寂灭→苏醒→悬念", size=10, color=RGBColor(0x88,0x88,0x88))
add_sep()

# H5
add_h3("【H5】近景 · 超新星+抢刀")
add("男生单手撑地——怪物将残余能量汇聚至左侧光核。光核亮度达到全片最高——刺眼白蓝光将深坑照成白昼——碎石被吸离地面漂浮。刀在五步外。他爬起——碎石打脸——扑向刀——手指触到刀柄。身后雾墙在超新星强光下被映成半透明的白色幕墙。")
add_note("✅ 4事件：汇聚→爬起→扑刀→触柄。规则9：垂死超新星。雾墙重申。")

# H6
add_h3("【H6】低角度正面后拉 · 终结一击")
add("握刀转身——左脚蹬地——碎石环形炸裂——直线冲刺。刀在身后拖出火星轨迹——三步后全力跃起——双手握刀过顶——刀尖刺入左侧光核。光核在刀尖触碰瞬间——先是一瞬绝对安静——然后白光吞没一切。画面全白。")
add_note("✅ 4事件：握刀→冲刺→跃起→刺入。规则3：蹬地→直线冲刺→跃起→下刺。")

# H7a
add_h3("【H